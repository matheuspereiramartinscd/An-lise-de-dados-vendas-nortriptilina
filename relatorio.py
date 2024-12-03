import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL

# Função para melhorar a aparência dos gráficos
def melhorar_grafico():
    plt.style.use('seaborn-v0_8-darkgrid')  # Usando estilo de gráfico válido
    plt.rcParams.update({'figure.figsize': (8, 6), 'font.size': 12})  # Tamanho da fonte e figura

# Função para formatar valores como moeda
def formatar_moeda(valor):
    try:
        return f'R$ {valor:,.2f}'  # Tenta formatar como moeda
    except ValueError:
        return str(valor)  # Se falhar, retorna o valor como string

# Função para gerar o relatório em Word
def gerar_relatorio_word(dados):
    doc = Document()
    doc.add_heading(f'Relatório de Análise de Vendas - {datetime.now().strftime("%d/%m/%Y")}', 0)

    # --- 2. Quantidade Total de Produtos Vendidos ---
    qtd_total_produtos = dados['QTD_UNIDADE_FARMACOTECNICA'].sum() if 'QTD_UNIDADE_FARMACOTECNICA' in dados.columns else 0
    doc.add_heading('Quantidade Total de Produtos Vendidos', level=1)
    doc.add_paragraph(f'{qtd_total_produtos}')

    # --- Métrica: Total e mês com mais vendas em 2020 e 2021 ---
    dados['MÊS_VENDA'] = pd.to_datetime(dados['MÊS_VENDA'], format='%B', errors='coerce')  # Corrigir para mês por nome
    dados['ANO_VENDA'] = dados['ANO_VENDA'].astype(str)

    vendas_por_ano_mes = dados.groupby(['ANO_VENDA', dados['MÊS_VENDA'].dt.month])['QTD_UNIDADE_FARMACOTECNICA'].sum()
    vendas_2020 = vendas_por_ano_mes.loc[2020].sum() if 2020 in vendas_por_ano_mes.index else 0
    vendas_2021 = vendas_por_ano_mes.loc[2021].sum() if 2021 in vendas_por_ano_mes.index else 0

    mes_top_2020 = vendas_por_ano_mes.loc[2020].idxmax() if 2020 in vendas_por_ano_mes.index else ("", 0)
    vendas_top_2020 = vendas_por_ano_mes.loc[2020].max() if 2020 in vendas_por_ano_mes.index else 0

    mes_top_2021 = vendas_por_ano_mes.loc[2021].idxmax() if 2021 in vendas_por_ano_mes.index else ("", 0)
    vendas_top_2021 = vendas_por_ano_mes.loc[2021].max() if 2021 in vendas_por_ano_mes.index else 0

    doc.add_heading('Vendas Totais e Mês com Mais Vendas (2020 e 2021)', level=1)
    doc.add_paragraph(f"Total de vendas em 2020: {vendas_2020:,.0f}")
    doc.add_paragraph(f"Mês com mais vendas em 2020: {mes_top_2020} (Vendas: {vendas_top_2020:,.0f})")
    doc.add_paragraph(f"Total de vendas em 2021: {vendas_2021:,.0f}")
    doc.add_paragraph(f"Mês com mais vendas em 2021: {mes_top_2021} (Vendas: {vendas_top_2021:,.0f})")

    # --- Tendência de Vendas ao Longo do Tempo ---
    vendas_mensais = dados.groupby(dados['MÊS_VENDA'].dt.to_period('M'))['QTD_UNIDADE_FARMACOTECNICA'].sum()
    plt.figure()
    sns.lineplot(x=vendas_mensais.index.astype(str), y=vendas_mensais.values, color='green')
    plt.title('Tendência de Vendas por Mês', fontsize=14)
    plt.xlabel('Mês')
    plt.ylabel('Quantidade Vendida')
    plt.xticks(rotation=45)
    vendas_grafico = 'vendas_mensais.png'
    plt.savefig(vendas_grafico)
    doc.add_heading('Tendência de Vendas ao Longo do Tempo', level=1)
    doc.add_paragraph("Análise da tendência de vendas mês a mês.")
    doc.add_picture(vendas_grafico, width=Inches(5.5))

    # --- 3. Média de Vendas por Mês ---
    vendas_por_mes = dados.groupby(dados['MÊS_VENDA'].dt.to_period('M'))['QTD_UNIDADE_FARMACOTECNICA'].sum()
    media_vendas_por_mes = vendas_por_mes.mean()
    doc.add_heading('Média de Vendas por Mês', level=1)
    doc.add_paragraph(f'{media_vendas_por_mes:,.0f}')

    # --- 4. Mediana de Vendas por Mês ---
    mediana_vendas_por_mes = vendas_por_mes.median()
    doc.add_heading('Mediana de Vendas por Mês', level=1)
    doc.add_paragraph(f'{mediana_vendas_por_mes:,.0f}')

    # --- 5. Desvio Padrão de Vendas por Mês ---
    desvio_padrao_vendas_por_mes = vendas_por_mes.std()
    doc.add_heading('Desvio Padrão de Vendas por Mês', level=1)
    doc.add_paragraph(f'{desvio_padrao_vendas_por_mes:,.0f}')

    # --- 6. Vendas por Sexo ---
    vendas_por_sexo = dados.groupby('GÊNERO')['QTD_UNIDADE_FARMACOTECNICA'].sum()
    doc.add_heading('Vendas por Sexo', level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sexo'
    hdr_cells[1].text = 'Quantidade Vendida'
    for sexo, qtd in vendas_por_sexo.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(sexo)
        row_cells[1].text = f'{qtd:,.0f}'

    # --- 7. Vendas por Estado (UF) ---
    vendas_por_estado = dados.groupby('UF_VENDA')['QTD_UNIDADE_FARMACOTECNICA'].sum()
    doc.add_heading('Vendas por Estado', level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Estado (UF)'
    hdr_cells[1].text = 'Quantidade Vendida'
    for estado, qtd in vendas_por_estado.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(estado)
        row_cells[1].text = f'{qtd:,.0f}'

    # --- 8. Vendas por Município ---
    vendas_por_municipio = dados.groupby('MUNICIPIO_VENDA')['QTD_UNIDADE_FARMACOTECNICA'].sum()
    doc.add_heading('Vendas por Município', level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Município'
    hdr_cells[1].text = 'Quantidade Vendida'
    for municipio, qtd in vendas_por_municipio.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(municipio)
        row_cells[1].text = f'{qtd:,.0f}'

    # --- 9. Distribuição de Idade dos Clientes ---
    dados_filtrados = dados[dados['IDADE'] > 10]  # Excluir idades 0
    melhorar_grafico()
    plt.figure()
    sns.histplot(dados_filtrados['IDADE'], kde=True, color='skyblue')
    plt.title('Distribuição de Idade dos Clientes (Sem Idade 0)', fontsize=14)
    plt.xlabel('Idade')
    plt.ylabel('Frequência')
    idade_grafico = 'idade_dist.png'
    plt.savefig(idade_grafico)
    doc.add_heading('Distribuição de Idade dos Clientes', level=1)
    doc.add_paragraph("Distribuição de idade dos clientes com base nas vendas realizadas.")
    doc.add_picture(idade_grafico, width=Inches(5.5))

    # --- 11. Correlação entre Idade e Quantidade Vendida ---
    plt.figure()
    sns.scatterplot(x=dados_filtrados['IDADE'], y=dados_filtrados['QTD_UNIDADE_FARMACOTECNICA'])
    plt.title('Correlação entre Idade e Quantidade Vendida', fontsize=14)
    plt.xlabel('Idade')
    plt.ylabel('Quantidade Vendida')
    idade_vendas_grafico = 'idade_vendas_corr.png'
    plt.savefig(idade_vendas_grafico)
    doc.add_heading('Correlação entre Idade e Quantidade Vendida', level=1)
    doc.add_paragraph("Análise de correlação entre idade do cliente e a quantidade de produtos vendidos.")
    doc.add_picture(idade_vendas_grafico, width=Inches(5.5))

    # Salvar o documento
    doc.save('relatorio_vendas_analise.docx')

# Exemplo de uso
# dados = pd.read_csv('vendas.csv')
# gerar_relatorio_word(dados)
